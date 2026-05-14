"""Build the final Word report from scratch using python-docx."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = "M11415803_Le_Trung_Kien_report.docx"


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, "2F5496")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm:
        for col, width in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(width)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.italic = True
    return p


def main():
    doc = Document()

    # Page setup: A4, 1 inch margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Link Prediction on a Directed Social Network")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HW1.2 - Final Report")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Author block
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("Le Trung Kien").bold = True
    author.add_run("    |    ")
    author.add_run("Student ID: M11415803").bold = True

    doc.add_paragraph()  # blank line

    # ── Task ────────────────────────────────────────────────────────────────
    doc.add_heading("Task", level=1)
    add_para(doc,
             "Predict hidden directed edges (Node1 → Node2) in a social-network graph "
             "reconstructed from approximately 186,000 observed training edges. "
             "Approximately 33,000 hidden edges are present in the test set. "
             "The evaluation metric is ROC AUC (Area Under the Receiver Operating "
             "Characteristic Curve).")

    # ── 1. Approach Summary ─────────────────────────────────────────────────
    doc.add_heading("1. Approach Summary", level=1)
    add_para(doc,
             "Two complementary supervised models were trained, and their outputs were "
             "blended in rank space to produce the final prediction.")
    add_table(doc,
              headers=["Component", "Role"],
              rows=[
                  ["Model A (main_1.py)",
                   "Compact hand-crafted graph features with LightGBM classifier."],
                  ["Model B (main_2.py)",
                   "Extended graph features plus 256-dim SVD embeddings with LightGBM."],
                  ["Rank blend (make_blends.py)",
                   "Weighted average of percentile ranks: 0.8 × Model A + 0.2 × Model B."],
              ],
              col_widths_cm=[5.0, 11.0])
    add_para(doc,
             "The blend was chosen because the two models exhibit moderate disagreement "
             "(Spearman ρ = 0.7226 between their predictions), indicating that combining "
             "them captures orthogonal signal not present in either model alone.")

    # ── 2. Method and Algorithm Design ──────────────────────────────────────
    doc.add_heading("2. Method and Algorithm Design", level=1)

    doc.add_heading("2.1 Algorithm Structure and Preprocessing Rationale", level=2)
    add_para(doc,
             "The link-prediction task is reformulated as a supervised binary "
             "classification problem. The pipeline proceeds in three stages: "
             "(i) build a graph from the training edges, (ii) extract features per "
             "candidate (u, v) pair, and (iii) train a classifier on observed edges "
             "plus synthetic negatives. Two such pipelines are run in parallel "
             "(Models A and B, Sections 2.2 and 2.3), and their predictions are "
             "combined in a final blending step (Section 2.4).")
    add_para(doc,
             "Every preprocessing step is listed below together with the reason for "
             "performing it.")
    add_table(doc,
              headers=["Preprocessing step", "Description and rationale"],
              rows=[
                  ["1. Graph construction",
                   "Train edges (Node1, Node2) are loaded into a NetworkX DiGraph; "
                   "the undirected projection is also built. "
                   "Reason: most local similarity metrics (Jaccard, Adamic-Adar) "
                   "are conventionally defined on undirected graphs, while "
                   "reciprocity and direction-aware metrics need the directed view."],
                  ["2. Adjacency caching",
                   "Successor sets, predecessor sets, and in/out degrees are cached "
                   "as Python dicts before feature extraction begins (both models). "
                   "Reason: NetworkX's per-call successors() is O(degree). With "
                   "~370 K (u, v) pairs and many neighbour lookups per pair, "
                   "caching reduces each lookup to O(1) and brings feature "
                   "extraction from minutes to seconds."],
                  ["3. Negative sampling",
                   "Random non-existent edges are sampled in equal count to the "
                   "positive edges, producing balanced training data. "
                   "Reason: link prediction has only positive labels; a "
                   "discriminative classifier needs negatives. Random sampling is "
                   "the standard baseline. Limitation: random negatives are easy "
                   "to separate from positives, which inflates CV AUC; harder "
                   "alternatives are discussed in Section 5.1."],
                  ["4. Pair-wise feature engineering",
                   "Each (u, v) pair is mapped to a fixed-length numeric vector "
                   "(19 dimensions for Model A, ~290 for Model B). "
                   "Reason: gradient-boosted trees require tabular inputs. The "
                   "chosen features (common neighbours, Adamic-Adar, preferential "
                   "attachment, etc.) have a long track record on social-network "
                   "link-prediction benchmarks."],
                  ["5. SVD on the adjacency matrix (Model B only)",
                   "The directed adjacency A is factorised by truncated SVD into "
                   "256-dim source and destination embeddings per node. "
                   "Reason: hand-crafted features capture local topology but miss "
                   "latent role/community structure. SVD is a cheap, deterministic "
                   "embedding that complements local features. The Hadamard "
                   "product of source and destination embeddings is a standard "
                   "pair-wise representation for link prediction."],
                  ["6. Pre-computed global node scores (Model B only)",
                   "PageRank, HITS hub/authority, and Louvain communities are "
                   "computed once and cached per node. "
                   "Reason: these scores express global importance and group "
                   "membership - information orthogonal to local neighbourhood "
                   "features - and are too expensive to recompute per (u, v) pair."],
                  ["7. LightGBM classifier",
                   "Gradient-boosted decision trees are used for both models. "
                   "Reason: LightGBM handles features with very different scales "
                   "(degrees in 1-1000 vs. similarities in [0, 1]) without "
                   "normalisation, captures non-linear interactions automatically, "
                   "trains quickly on ~370 K rows, and provides built-in early "
                   "stopping and feature-importance reporting."],
                  ["8. Stratified 5-fold cross-validation with fold-averaged test predictions",
                   "Both models train 5 fold-models on stratified splits and average "
                   "their test predictions to produce the final submission. "
                   "Reason: stratification preserves the 1:1 positive/negative "
                   "ratio in every fold; 5 folds balance variance of the CV "
                   "estimate against retraining cost. Averaging the 5 fold-trained "
                   "models acts as a small ensemble that reduces prediction "
                   "variance compared to a single classifier refit on all data."],
                  ["9. Rank-space prediction blending",
                   "Each model's probability vector is replaced by its percentile "
                   "ranks before averaging. "
                   "Reason: the two models output probabilities on very different "
                   "scales (mean 0.111 vs. 0.409). Naive probability averaging "
                   "would let the higher-mean model dominate regardless of the "
                   "weight. Percentile ranks normalise both models to a uniform "
                   "(0, 1] distribution, preserving only relative ordering - "
                   "which is what ROC AUC measures."],
              ],
              col_widths_cm=[5.0, 11.0])
    add_para(doc,
             "Steps 1-4 and 7-8 are shared by both models. Steps 5-6 "
             "(SVD embeddings, global node scores) are Model B only. "
             "Step 9 combines the two models' predictions into the final submission.")

    doc.add_heading("2.2 Model A - Hand-Crafted Graph Features", level=2)
    add_para(doc,
             "For each candidate pair (u, v), a 19-dimensional feature vector is computed "
             "from the directed graph G and its undirected projection. The features fall "
             "into the following groups:")
    add_table(doc,
              headers=["Group", "Features"],
              rows=[
                  ["Degrees",
                   "out_deg(u), in_deg(u), out_deg(v), in_deg(v)"],
                  ["Common neighbours",
                   "Four directed variants: |out_u ∩ out_v|, |out_u ∩ in_v|, "
                   "|in_u ∩ out_v|, |in_u ∩ in_v|; plus undirected count."],
                  ["Similarity",
                   "Jaccard coefficient (undirected), Adamic-Adar index"],
                  ["Preferential attachment",
                   "out_deg(u) × in_deg(v); |N_u| × |N_v|"],
                  ["Reciprocity",
                   "v→u edge indicator; u ∈ out(v) indicator"],
                  ["Path / structure",
                   "Length-2 path count; hub ratios for u and v; "
                   "Katz-style proxy: cn_out_out / (out_deg(u) + 1)"],
              ],
              col_widths_cm=[4.5, 11.5])

    add_para(doc, "Training procedure:")
    add_bullet(doc,
               "Negative samples: random non-edges drawn equal in count to "
               "positives (random seed = 42).")
    add_bullet(doc,
               "Classifier: LightGBM with n_estimators = 800, learning_rate = 0.05, "
               "num_leaves = 63, max_depth = 8, subsample = 0.8, "
               "colsample_bytree = 0.8, reg_alpha = 0.1, reg_lambda = 1.0.")
    add_bullet(doc, "Early stopping: patience = 50 rounds.")
    add_bullet(doc,
               "Cross-validation: stratified 5-fold; final test predictions are "
               "the per-fold average.")

    doc.add_heading("2.3 Model B - Graph Features and SVD Embeddings", level=2)
    add_para(doc,
             "Model B adds two layers of richness on top of Model A's design.")
    add_para(doc,
             "(i) Twenty-seven graph features, including everything in Model A plus "
             "three additional groups:")
    add_bullet(doc,
               "Global node scores: PageRank, HITS hub score, HITS authority score "
               "for both u and v.")
    add_bullet(doc,
               "Community membership: Louvain communities computed on the undirected "
               "graph; same_community indicator.")
    add_bullet(doc,
               "Resource-allocation index, triadic-closure ratio, follow-ratios "
               "(out_deg / (in_deg + 1)).")

    add_para(doc,
             "(ii) 256-dimensional truncated SVD embeddings of the binary adjacency "
             "matrix A. Letting A be approximated by U·Σ·Vᵀ (top-256 singular values), "
             "the source embedding for node u is U[u]·√Σ and the destination embedding "
             "for node v is Vᵀ[v]·√Σ.")
    add_para(doc,
             "For each pair, the embedding-derived features are: dot product, "
             "L2 distance, cosine similarity, source/destination norms, summary "
             "statistics of the Hadamard product (mean, standard deviation, "
             "absolute mean, absolute max), plus the full 256-dimensional Hadamard "
             "product itself.")
    add_para(doc,
             "Total feature dimension is approximately 290. Classifier: LightGBM "
             "with n_estimators = 1000, learning_rate = 0.03, num_leaves = 255. "
             "Stratified 5-fold cross-validation.")

    doc.add_heading("2.4 Final Ensemble - Rank-Space Blending", level=2)
    add_para(doc,
             "The two models exhibit very different probability calibrations:")
    add_table(doc,
              headers=["Model", "Mean predicted probability", "Median"],
              rows=[
                  ["Model A", "0.111", "0.0004"],
                  ["Model B", "0.409", "0.0056"],
              ],
              col_widths_cm=[4.0, 6.0, 4.0])
    add_para(doc,
             "Naive probability-space averaging would be dominated by Model B's "
             "inflated scale. To remove the scale mismatch, predictions are first "
             "converted to percentile ranks in [0, 1], then averaged:")
    add_code_line(doc,
                  "final_score(pair) = 0.8 × rank_A(pair) + 0.2 × rank_B(pair)")
    add_para(doc,
             "The weight 0.8 was selected by probing four candidates (50/50, 70/30, "
             "80/20, 90/10) on the public leaderboard; 80/20 gave the highest score "
             "(see Section 3). Because rank-AUC is a step function of the blend "
             "weight (it changes only at points where two pairs swap rank order), "
             "the curve is not smooth and the optimum has to be found empirically.")

    # ── 3. Results ──────────────────────────────────────────────────────────
    doc.add_heading("3. Results", level=1)
    add_table(doc,
              headers=["Submission", "CV AUC", "Public LB AUC"],
              rows=[
                  ["Model A (main_1.py)", "0.99997", "0.85428"],
                  ["Model B (main_2.py)", "0.99940", "0.78773"],
                  ["Rank blend 50/50", "-", "0.84212"],
                  ["Rank blend 70/30", "-", "0.85292"],
                  ["Rank blend 90/10", "-", "0.85062"],
                  ["Rank blend 80/20  (final pick #1)", "-", "0.85505"],
                  ["Model A standalone (final pick #2)", "-", "0.85428"],
              ],
              col_widths_cm=[8.0, 4.0, 4.0])
    add_para(doc,
             "Final submissions. The two final picks were chosen for diversity. The "
             "80/20 rank blend captures both models' signals and gives the highest "
             "public-LB score. Pure Model A acts as an anchor in case the blend "
             "overfits the public leaderboard.")
    p = doc.add_paragraph()
    run = p.add_run("Note on the CV-LB gap. ")
    run.bold = True
    p.add_run(
        "Both models scored near 1.0 in cross-validation but markedly lower on the "
        "leaderboard. The cause is the negative-sampling distribution mismatch: random "
        "non-edges are easy to distinguish from real edges, while the actual hidden "
        "edges lie in much harder regions of the graph (plausible non-observed "
        "connections). Model A's simpler feature set generalizes better than Model B's "
        "richer one (smaller CV-LB gap), which is why Model A dominates the blend."
    )

    # ── 4. How to Run ───────────────────────────────────────────────────────
    doc.add_heading("4. How to Run", level=1)

    doc.add_heading("4.1 Requirements", level=2)
    add_bullet(doc, "Python ≥ 3.12")
    add_bullet(doc, "Package manager: uv  (https://docs.astral.sh/uv/)")
    add_bullet(doc,
               "Dependencies (declared in pyproject.toml): lightgbm, networkx, "
               "numpy, pandas, scikit-learn, scipy, tqdm")
    add_para(doc, "Install dependencies:")
    add_code_line(doc, "uv sync")

    doc.add_heading("4.2 Input Data", level=2)
    add_para(doc, "Place the following files in the working directory:")
    add_bullet(doc, "train.csv - known edges  (columns: Node1, Node2)")
    add_bullet(doc, "test.csv - candidate pairs to score  (columns: ID, Node1, Node2)")
    add_bullet(doc, "sample_submission.csv - submission template  (columns: ID, Label)")

    doc.add_heading("4.3 Step 1 - Train Model A", level=2)
    add_para(doc,
             "Run the Model A pipeline. It produces submission.csv, which is "
             "renamed to keep separate from the Model B output:")
    add_code_line(doc, "uv run python main_1.py")
    add_code_line(doc, "mv submission.csv  M11415803_Le_Trung_Kien.csv")

    doc.add_heading("4.4 Step 2 - Train Model B", level=2)
    add_para(doc, "Run the Model B pipeline and rename the output:")
    add_code_line(doc, "uv run python main_2.py")
    add_code_line(doc, "mv submission.csv  M11415803_Le_Trung_Kien_0505_opus5fold.csv")

    doc.add_heading("4.5 Step 3 - Generate the Rank-Space Blend", level=2)
    add_para(doc,
             "Reads the two CSVs above and writes all blend variants:")
    add_code_line(doc, "uv run python make_blends.py")
    add_para(doc,
             "The selected final file is M11415803_Le_Trung_Kien_blend_rank_80_20.csv.")

    doc.add_heading("4.6 Reproducibility Notes", level=2)
    add_bullet(doc,
               "All RNG seeds are fixed (random_state = 42 in LightGBM, seeded "
               "numpy and random for negative sampling).")
    add_bullet(doc,
               "scipy.sparse.linalg.svds does not expose a deterministic seed for "
               "its ARPACK initialization, so Model B's outputs may differ by "
               "approximately 10⁻⁴ in absolute probability across runs. Rank order, "
               "and therefore AUC, is essentially unaffected.")

    # ── 5. Future Improvements ──────────────────────────────────────────────
    doc.add_heading("5. Future Improvements", level=1)
    add_para(doc,
             "The current pipeline has clear weaknesses, observed during model "
             "development. The following improvements are ordered by expected "
             "impact on private-leaderboard performance.")

    doc.add_heading("5.1 Hard Negative Sampling (Highest Priority)", level=2)
    add_para(doc,
             "The dominant root cause of the large CV-LB gap (~0.15 AUC for both "
             "models) is random negative sampling. Random non-edges in a sparse "
             "social network are trivially distinguishable from real edges (very "
             "low degree, few common neighbours), while the actual hidden test "
             "edges are plausible non-observed connections.")
    add_para(doc, "Replace the sample_random_negatives function with at least one of:")
    add_bullet(doc,
               "2-hop sampling: pairs (u, v) where v is reachable from u in exactly "
               "2 hops but no direct edge exists.")
    add_bullet(doc,
               "Preferential-attachment-weighted sampling: draw u and v with "
               "probability proportional to their degrees.")
    add_bullet(doc,
               "Within-community sampling: draw pairs from the same Louvain "
               "community (homophily-aware negatives).")
    add_para(doc,
             "Expected impact: lower CV AUC (closer to LB) and higher absolute "
             "LB AUC, plausibly +0.02 to +0.05.")

    doc.add_heading("5.2 Better Embedding Methods", level=2)
    add_para(doc,
             "Truncated SVD on the adjacency matrix encodes \"what is already "
             "connected\", which biases the model toward visible structure. "
             "Stronger alternatives:")
    add_bullet(doc,
               "Node2vec or DeepWalk: random-walk-based embeddings that capture "
               "both local and global structure.")
    add_bullet(doc,
               "LINE: optimised specifically for first- and second-order proximity.")
    add_bullet(doc,
               "Graph autoencoders (GAE / VGAE): trained directly with a "
               "link-prediction objective.")
    add_bullet(doc,
               "GraphSAGE or other GNNs: learn an asymmetric scoring function "
               "for (u, v) pairs end-to-end.")

    doc.add_heading("5.3 Probability Calibration", level=2)
    add_para(doc,
             "Model A and Model B differ in calibration by approximately 4× in "
             "mean predicted probability. Fitting isotonic regression on a "
             "held-out fold for each model would make probability-space averaging "
             "viable, removing the need for the rank-space workaround and "
             "potentially enabling further blends with additional models.")

    doc.add_heading("5.4 Stacking", level=2)
    add_para(doc,
             "The two models could feed their out-of-fold predictions into a "
             "meta-learner (e.g. logistic regression or a small LightGBM) along "
             "with a few hand-picked graph statistics, instead of being averaged "
             "with a fixed weight. This typically gives a small but reliable "
             "improvement (+0.002 to +0.005 AUC in similar problems).")

    doc.add_heading("5.5 Hyperparameter Tuning", level=2)
    add_para(doc,
             "Both models use intuitively-chosen hyperparameters. Bayesian "
             "optimisation (Optuna or similar) over a CV objective with hard "
             "negatives would likely yield further gains, particularly around "
             "the larger num_leaves in Model B and the regularisation coefficients.")

    doc.add_heading("5.6 Cost-Benefit Summary", level=2)
    add_table(doc,
              headers=["Improvement", "Code effort", "Expected gain (LB AUC)"],
              rows=[
                  ["Hard negative sampling", "10-30 lines", "+0.02 to +0.05"],
                  ["Probability calibration", "20-40 lines", "+0.001 to +0.005"],
                  ["Stacking meta-learner", "50-100 lines", "+0.002 to +0.005"],
                  ["Hyperparameter tuning (Optuna)",
                   "30-50 lines, several CPU hours", "+0.001 to +0.01"],
                  ["Node2vec / DeepWalk embeddings",
                   "Library call, ~1 hour CPU", "+0.005 to +0.02"],
                  ["GraphSAGE / GAE",
                   "100+ lines, 1-2 hours GPU", "+0.01 to +0.03"],
              ],
              col_widths_cm=[6.0, 5.5, 4.5])
    add_para(doc,
             "Hard-negative sampling alone is the highest-impact, lowest-effort "
             "change and should be attempted first in any future iteration.")

    # ── 6. File Manifest ────────────────────────────────────────────────────
    doc.add_heading("6. File Manifest", level=1)
    add_table(doc,
              headers=["File", "Description"],
              rows=[
                  ["main_1.py",
                   "Model A pipeline (features + 5-fold LightGBM)."],
                  ["main_2.py",
                   "Model B pipeline (features + SVD + LightGBM)."],
                  ["make_blends.py",
                   "Computes rank-space and probability-space blends."],
                  ["M11415803_Le_Trung_Kien.csv",
                   "Model A submission - final pick #2."],
                  ["M11415803_Le_Trung_Kien_blend_rank_80_20.csv",
                   "Final blend - final pick #1."],
                  ["M11415803_Le_Trung_Kien_report.docx",
                   "This report."],
              ],
              col_widths_cm=[7.5, 8.5])

    doc.save(OUTPUT)
    print(f"Saved -> {OUTPUT}")


if __name__ == "__main__":
    main()
